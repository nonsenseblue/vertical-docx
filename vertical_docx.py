#!/usr/bin/env python3
"""
vertical-docx — Markdown to vertical-writing docx converter
"""

import argparse
import os
import sys

from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BLACK = RGBColor(0, 0, 0)
GREY = RGBColor(180, 180, 180)
QUOTE_GREY = RGBColor(102, 102, 102)

W_NS = nsdecls("w")


def half_to_full(text):
    result = []
    for ch in text:
        cp = ord(ch)
        if 0x21 <= cp <= 0x7E:
            result.append(chr(cp + 0xFEE0))
        elif ch == ' ':
            result.append('　')
        else:
            result.append(ch)
    return ''.join(result)


class VerticalDocxBuilder:

    def __init__(self, cfg):
        self.cfg = cfg
        self.font = cfg.font
        self.size = Pt(cfg.font_size)
        self.doc = Document()
        self._init_style()
        self._init_layout()
        if not cfg.no_page_numbers:
            self._add_page_numbers()

    # -- setup ---------------------------------------------------------------

    def _init_style(self):
        style = self.doc.styles['Normal']
        style.font.name = self.font
        style.font.size = self.size
        style.font.color.rgb = BLACK
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        self._set_east_asian_font(style.element)

    def _init_layout(self):
        cfg = self.cfg
        for section in self.doc.sections:
            if cfg.landscape:
                section.page_width, section.page_height = Mm(297), Mm(210)
                section.orientation = 1
                printable_w = 297 - cfg.margin_left - cfg.margin_right
            else:
                section.page_width, section.page_height = Mm(210), Mm(297)
                printable_w = 210 - cfg.margin_left - cfg.margin_right

            section.top_margin = Mm(cfg.margin_top)
            section.bottom_margin = Mm(cfg.margin_bottom)
            section.left_margin = Mm(cfg.margin_left)
            section.right_margin = Mm(cfg.margin_right)

            sect_pr = section._sectPr
            sect_pr.append(parse_xml(
                f'<w:textDirection {W_NS} w:val="tbRl"/>'
            ))

            grid = sect_pr.find(qn('w:docGrid'))
            if grid is None:
                grid = parse_xml(f'<w:docGrid {W_NS}/>')
                sect_pr.append(grid)

            pitch = cfg.line_pitch or int(printable_w * 2.8346 / cfg.lines * 20)
            grid.set(qn('w:type'), 'linesAndChars')
            grid.set(qn('w:linePitch'), str(pitch))
            grid.set(qn('w:charSpace'), str(cfg.char_space))

    def _add_page_numbers(self):
        for section in self.doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            p = footer.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._add_run(p, xml=f'<w:fldChar {W_NS} w:fldCharType="begin"/>', size=Pt(9))
            self._add_run(p, xml=f'<w:instrText {W_NS} xml:space="preserve"> PAGE </w:instrText>', size=Pt(9))
            self._add_run(p, xml=f'<w:fldChar {W_NS} w:fldCharType="end"/>', size=Pt(9))

    # -- helpers -------------------------------------------------------------

    def _set_east_asian_font(self, element):
        rPr = element.find(qn('w:rPr'))
        if rPr is None:
            return
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {W_NS}/>')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), self.font)

    def _style_run(self, run, size=None, bold=False, color=None):
        run.font.name = self.font
        run.font.color.rgb = color or BLACK
        run.font.bold = bold
        if size:
            run.font.size = size
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {W_NS}/>')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), self.font)
        return run

    def _add_run(self, paragraph, text='', size=None, bold=False, color=None, xml=None):
        run = paragraph.add_run(text)
        self._style_run(run, size=size, bold=bold, color=color)
        if xml:
            run._element.append(parse_xml(xml))
        return run

    def _add_paragraph(self, text='', size=None, align=None,
                       space_before=None, space_after=None,
                       page_break_before=False, color=None):
        p = self.doc.add_paragraph()
        if align:
            p.alignment = align
        if page_break_before:
            p.paragraph_format.page_break_before = True
        if space_before is not None:
            p.paragraph_format.space_before = space_before
        if space_after is not None:
            p.paragraph_format.space_after = space_after
        self._add_run(p, text, size=size or self.size, color=color)
        return p

    def _empty_lines(self, n):
        for _ in range(n):
            self._add_paragraph()

    # -- content blocks ------------------------------------------------------

    def add_title_page(self, title, author=None):
        self._empty_lines(8)
        self._add_paragraph(title, size=Pt(22), align=WD_ALIGN_PARAGRAPH.CENTER)
        self._empty_lines(3)
        if author:
            self._add_paragraph(
                half_to_full(author), size=Pt(14),
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
        self.doc.add_page_break()

    def add_chapter_page(self, title):
        self._add_paragraph(page_break_before=True)
        self._empty_lines(5)
        self._add_paragraph(title, size=Pt(16), align=WD_ALIGN_PARAGRAPH.CENTER)
        self.doc.add_page_break()

    def add_section_header(self, title):
        self._add_paragraph(
            title, align=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=Pt(24), space_after=Pt(18),
        )

    def add_body(self, text):
        self._add_paragraph(text, space_before=Pt(0), space_after=Pt(0))

    def add_blockquote(self, text):
        self._add_paragraph(text, size=Pt(9), color=QUOTE_GREY)

    def add_scene_break(self):
        self._add_paragraph(
            '＊', align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY,
        )

    def add_synopsis_heading(self, title):
        self._add_paragraph(
            title, size=Pt(18),
            align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(24),
        )

    # -- markdown processing -------------------------------------------------

    def process_md(self, path, is_synopsis=False):
        with open(path, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        in_code = False
        first_h1 = True

        for line in lines:
            s = line.rstrip('\n')

            if s.startswith('```'):
                in_code = not in_code
                continue
            if in_code:
                self.add_body(s)
                continue

            clean = s.replace('**', '').replace('`', '')

            if s.startswith('# ') and not s.startswith('## '):
                title = half_to_full(clean[2:])
                if first_h1 and not is_synopsis:
                    self.add_title_page(
                        self.cfg.title or title, self.cfg.author,
                    )
                    first_h1 = False
                elif is_synopsis:
                    self.add_synopsis_heading(title)
                    first_h1 = False
                else:
                    self.add_chapter_page(title)

            elif s.startswith('## '):
                self.add_chapter_page(half_to_full(clean[3:]))

            elif s.startswith('### '):
                self.add_section_header(half_to_full(clean[4:]))

            elif s.startswith('> '):
                self.add_blockquote(half_to_full(clean[2:]))

            elif s.strip() in ('---', '-----'):
                self.add_scene_break()

            elif s.strip():
                self.add_body(half_to_full(clean))

    # -- save ----------------------------------------------------------------

    def save(self, path):
        self.doc.save(path)
        return os.path.getsize(path)


def main():
    parser = argparse.ArgumentParser(
        description='Markdown to vertical-writing docx converter',
    )
    parser.add_argument('input', help='Input markdown file')
    parser.add_argument('-o', '--output', help='Output docx path')
    parser.add_argument('--title', help='Override document title (default: first # heading)')
    parser.add_argument('--author', help='Author name for title page')
    parser.add_argument('--synopsis', help='Path to synopsis markdown (prepended)')
    parser.add_argument('--font', default='Yu Mincho', help='Font name (default: Yu Mincho)')
    parser.add_argument('--font-size', type=float, default=10.5, help='Font size in pt (default: 10.5)')
    parser.add_argument('--chars', type=int, default=40, help='Characters per line (default: 40)')
    parser.add_argument('--lines', type=int, default=30, help='Lines per page (default: 30)')
    parser.add_argument('--landscape', action='store_true', help='Use landscape orientation')
    parser.add_argument('--margin-top', type=float, default=30, help='Top margin in mm (default: 30)')
    parser.add_argument('--margin-bottom', type=float, default=25, help='Bottom margin in mm (default: 25)')
    parser.add_argument('--margin-left', type=float, default=25, help='Left margin in mm (default: 25)')
    parser.add_argument('--margin-right', type=float, default=25, help='Right margin in mm (default: 25)')
    parser.add_argument('--line-pitch', type=int, default=0, help='Override docGrid linePitch (advanced)')
    parser.add_argument('--char-space', type=int, default=0, help='Override docGrid charSpace (advanced)')
    parser.add_argument('--no-page-numbers', action='store_true', help='Disable page numbers')

    cfg = parser.parse_args()

    if not os.path.isfile(cfg.input):
        print(f'Error: {cfg.input} not found', file=sys.stderr)
        sys.exit(1)

    if not cfg.output:
        cfg.output = os.path.splitext(cfg.input)[0] + '.docx'

    builder = VerticalDocxBuilder(cfg)

    if cfg.synopsis:
        if not os.path.isfile(cfg.synopsis):
            print(f'Error: {cfg.synopsis} not found', file=sys.stderr)
            sys.exit(1)
        builder.process_md(cfg.synopsis, is_synopsis=True)
        builder.doc.add_page_break()

    builder.process_md(cfg.input)

    size = builder.save(cfg.output)
    print(f'{cfg.output} ({size / 1024:.0f} KB)')


if __name__ == '__main__':
    main()
